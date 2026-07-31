#!/usr/bin/env python3
"""Render selected synthetic fixtures into real files (MD / HTML / DOCX / PDF).

Reads `corpus.jsonl` and writes a representative subset as actual documents, so
the ingestion module (#4) and downstream checks can be exercised on real file
formats rather than text-in-JSON. Ground truth stays in `manifest.csv`, keyed by
id; this script also writes `files_index.csv` mapping each rendered file back to
its record and its ground-truth flags.

Formats:
    md, html   pure stdlib.
    docx       needs python-docx (the [docx] extra).
    pdf        LOCAL build step -- a Chromium-family browser prints the HTML.
               Set CHROMIUM to a browser binary, or have chrome/edge on PATH.
               CI never renders; it reads the committed files.

Examples
--------
    python3 render_fixtures.py --corpus tests/fixtures/synthetic --formats md html docx
    CHROMIUM="/Applications/Google Chrome.app/Contents/MacOS/Google Chrome" \\
        python3 render_fixtures.py --corpus tests/fixtures/synthetic --formats md html docx pdf
"""

import argparse
import csv
import html
import json
import os
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

CITATION_HEADING = {
    "grant_application": "References",
    "blog_post": "Sources",
    "think_tank_report": "Endnotes",
}

# The representative cases to render, per document type. Each entry is matched
# against the corpus by (label, defect-predicate); the first match is rendered.
GRANT_CASES = [
    "human",
    "ai_clean",
    "fabricated_citations",
    "wrong_paper",
    "overclaims",
    "budget_inflated",
    "missing_methods",
]
TEXT_CASES = ["ai_clean", "fabricated_citations", "wrong_paper", "overclaims"]
GROUND_TRUTH_FIELDS = [
    "ai_generated",
    "has_fabricated_citations",
    "has_mismatched_citations",
    "overclaims",
    "budget_inflated",
    "missing_methods",
]


def _matches(record, case):
    label, defect = record["label"], record["dimensions"]["defect"]
    if case in ("human", "ai_clean"):
        return label == case and defect == "none"
    return label == "slop" and defect in (case, "all")


def select_fixtures(records):
    by_type = {
        "grant_application": GRANT_CASES,
        "blog_post": TEXT_CASES,
        "think_tank_report": TEXT_CASES,
    }
    chosen, seen = [], set()
    for doc_type, cases in by_type.items():
        pool = [r for r in records if r["dimensions"]["document_type"] == doc_type]
        for case in cases:
            hit = next((r for r in pool if _matches(r, case) and r["id"] not in seen), None)
            if hit:
                chosen.append((case, hit))
                seen.add(hit["id"])
    return chosen


# --------------------------------------------------------------------------- #
# One normalized view of a document, rendered into each format
# --------------------------------------------------------------------------- #
def _blocks(record):
    title = record["title"]
    body = [(k.replace("_", " ").title(), v) for k, v in record["sections"].items()]
    heading = CITATION_HEADING.get(record["dimensions"]["document_type"], "References")
    cites = [f"{c['marker']} https://doi.org/{c['doi']}" for c in record["citations"]]
    return title, body, heading, cites


def render_md(record):
    title, body, heading, cites = _blocks(record)
    out = [f"# {title}", ""]
    for h, text in body:
        out.append(f"## {h}")
        out.append("")
        out.extend(line for line in text.split("\n"))
        out.append("")
    out.append(f"## {heading}")
    out.append("")
    out.extend(f"- {c}" for c in cites)
    return "\n".join(out) + "\n"


def render_html(record):
    title, body, heading, cites = _blocks(record)
    esc = html.escape
    parts = [
        "<!doctype html>",
        '<html lang="en"><head><meta charset="utf-8">',
        f"<title>{esc(title)}</title></head><body>",
        f"<h1>{esc(title)}</h1>",
    ]
    for h, text in body:
        parts.append(f"<h2>{esc(h)}</h2>")
        parts.extend(f"<p>{esc(line)}</p>" for line in text.split("\n") if line.strip())
    parts.append(f"<h2>{esc(heading)}</h2>")
    parts.append("<ol>")
    parts.extend(f"<li>{esc(c)}</li>" for c in cites)
    parts.append("</ol></body></html>")
    return "\n".join(parts) + "\n"


def render_docx(record, out_path):
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("docx format needs python-docx (pip install 'python-docx')") from e
    title, body, heading, cites = _blocks(record)
    doc = docx.Document()
    doc.add_heading(title, level=0)
    for h, text in body:
        doc.add_heading(h, level=1)
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
    doc.add_heading(heading, level=1)
    for c in cites:
        doc.add_paragraph(c, style="List Number")
    doc.save(str(out_path))


def find_browser():
    env = os.environ.get("CHROMIUM")
    if env and Path(env).exists():
        return env
    mac = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
    ]
    for path in mac:
        if Path(path).exists():
            return path
    for name in ("google-chrome", "chrome", "msedge", "chromium-browser", "chromium"):
        found = shutil.which(name)
        if found:
            return found
    return None


def render_pdf(record, out_path, browser, timeout=60):
    # Headless Chrome writes the PDF but often will not exit cleanly (macOS Chrome
    # lingers; on CI a crashpad child deadlocks the pipes -- the #47/#49 hang). So
    # don't wait for exit: poll until the PDF is written and its size stops
    # growing, then terminate the browser. The committed PDF is what we want.
    out_path = Path(out_path)
    if out_path.exists():
        out_path.unlink()
    with tempfile.TemporaryDirectory() as tmp:
        html_path = Path(tmp) / "doc.html"
        html_path.write_text(render_html(record), encoding="utf-8")
        proc = subprocess.Popen(
            [
                browser,
                "--headless=new",
                "--disable-gpu",
                "--no-sandbox",
                "--disable-crash-reporter",
                "--disable-breakpad",
                "--no-first-run",
                "--disable-dev-shm-usage",
                f"--user-data-dir={tmp}/profile",
                "--no-pdf-header-footer",
                f"--print-to-pdf={out_path}",
                html_path.as_uri(),
            ],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        try:
            deadline = time.monotonic() + timeout
            last, stable = -1, 0
            while time.monotonic() < deadline:
                if proc.poll() is not None:
                    break
                size = out_path.stat().st_size if out_path.exists() else 0
                stable = stable + 1 if size > 0 and size == last else 0
                last = size
                if stable >= 2:  # size unchanged across two polls -> write done
                    break
                time.sleep(0.4)
        finally:
            if proc.poll() is None:
                proc.terminate()
                try:
                    proc.wait(timeout=5)
                except subprocess.TimeoutExpired:
                    proc.kill()
    if not out_path.exists() or out_path.stat().st_size == 0:
        raise RuntimeError(f"browser produced no PDF at {out_path}")


# --------------------------------------------------------------------------- #
# Driver
# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser(description="Render synthetic fixtures into real files.")
    ap.add_argument("--corpus", default="tests/fixtures/synthetic", help="dir with corpus.jsonl")
    ap.add_argument("--out", default=None, help="output dir (default: <corpus>/files)")
    ap.add_argument(
        "--formats",
        nargs="+",
        default=["md", "html", "docx", "pdf"],
        choices=["md", "html", "docx", "pdf"],
    )
    args = ap.parse_args()

    corpus_dir = Path(args.corpus)
    out_dir = Path(args.out) if args.out else corpus_dir / "files"
    out_dir.mkdir(parents=True, exist_ok=True)
    records = [json.loads(line) for line in open(corpus_dir / "corpus.jsonl")]
    chosen = select_fixtures(records)

    browser = None
    if "pdf" in args.formats:
        browser = find_browser()
        if not browser:
            print("[warn] no Chromium-family browser found; skipping PDF")
            args.formats = [f for f in args.formats if f != "pdf"]

    index_rows = []
    for case, record in chosen:
        stem = f"{record['dimensions']['document_type']}__{case}"
        for fmt in args.formats:
            path = out_dir / f"{stem}.{fmt}"
            if fmt == "md":
                path.write_text(render_md(record), encoding="utf-8")
            elif fmt == "html":
                path.write_text(render_html(record), encoding="utf-8")
            elif fmt == "docx":
                render_docx(record, path)
            elif fmt == "pdf":
                render_pdf(record, path, browser)
            g = record["ground_truth"]
            index_rows.append(
                [
                    path.name,
                    record["id"],
                    record["dimensions"]["document_type"],
                    record["label"],
                    case,
                ]
                + [g[f] for f in GROUND_TRUTH_FIELDS]
            )

    index_path = out_dir / "files_index.csv"
    with open(index_path, "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["file", "corpus_id", "document_type", "label", "case", *GROUND_TRUTH_FIELDS])
        w.writerows(index_rows)

    print(f"Rendered {len(chosen)} fixtures x {len(args.formats)} formats -> {out_dir}/")
    print(f"  formats: {', '.join(args.formats)}")
    print(f"  index:   {index_path}")


if __name__ == "__main__":
    main()
